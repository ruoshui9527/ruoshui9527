from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import yaml
import json


class SwaggerToWord:
    def __init__(self, swagger_path, output_path):
        self.swagger_path = swagger_path
        self.output_path = output_path
        self.doc = Document()
        self.swagger_data = self._load_swagger()

    def _load_swagger(self):
        """加载Swagger文件"""
        with open(self.swagger_path, 'r', encoding='utf-8') as f:
            if self.swagger_path.endswith('.yaml') or self.swagger_path.endswith('.yml'):
                return yaml.safe_load(f)
            else:  # 假定为JSON格式
                return json.load(f)

    def _add_heading(self, text, level=1):
        """添加带样式的标题"""
        heading = self.doc.add_heading(level=level)
        run = heading.add_run(text)
        run.font.name = '微软雅黑'
        run.font.size = Pt(18 if level == 1 else 14)
        return heading

    def _add_paragraph(self, text, style=None):
        """添加段落"""
        p = self.doc.add_paragraph(text, style=style)
        p.paragraph_format.space_after = Pt(6)
        return p

    def _add_table(self, rows, cols):
        """创建表格并返回表格对象"""
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = 'Light Shading Accent 1'
        table.autofit = False
        table.allow_autofit = False
        for row in table.rows:
            for cell in row.cells:
                cell.width = Inches(2)
        return table

    def _process_parameters(self, parameters):
        """处理参数生成表格"""
        if not parameters:
            return
        table = self._add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '参数名'
        hdr_cells[1].text = '位置'
        hdr_cells[2].text = '类型'
        hdr_cells[3].text = '必需'
        hdr_cells[4].text = '描述'

        for param in parameters:
            row_cells = table.add_row().cells
            row_cells[0].text = param.get('name', '')
            row_cells[1].text = param.get('in', '')
            row_cells[2].text = param.get('type',
                                          param.get('schema', {}).get('type', ''))
            row_cells[3].text = '是' if param.get('required') else '否'
            row_cells[4].text = param.get('description', '')

    def _process_response(self, responses):
        """处理响应示例"""
        if not responses:
            return
        for code, response in responses.items():
            self._add_paragraph(f"响应状态: {code}", style='Heading3')
            if 'content' in response:
                content = response['content']
                for content_type, schema in content.items():
                    example = schema.get('example', '暂无示例')
                    self._add_paragraph(f"示例 ({content_type}):")
                    self.doc.add_paragraph(str(example), style='Intense Quote')

    def generate_doc(self):
        """生成Word文档主方法"""
        # 添加文档标题
        self._add_heading('API 文档', level=0)

        # 添加基本信息
        self._add_heading('基本信息', level=1)
        info = self.swagger_data.get('info', {})
        self._add_paragraph(f"标题: {info.get('title', '')}")
        self._add_paragraph(f"版本: {info.get('version', '')}")
        self._add_paragraph(f"描述: {info.get('description', '')}")

        # 处理每个API路径
        self._add_heading('API 列表', level=1)
        for path, methods in self.swagger_data.get('paths', {}).items():
            self._add_heading(path, level=2)
            for method, details in methods.items():
                if method.lower() not in ['get', 'post', 'put', 'delete']:
                    continue

                # 添加方法标题
                method_heading = self.doc.add_paragraph()
                method_heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                run = method_heading.add_run(method.upper())
                run.font.color.rgb = RGBColor(255, 0, 0) if method == 'post' else RGBColor(0, 0, 255)
                run.font.bold = True

                # 添加摘要和描述
                self._add_paragraph(f"摘要: {details.get('summary', '')}")
                self._add_paragraph(f"描述: {details.get('description', '')}")

                # 处理参数
                self._add_paragraph("请求参数:", style='Heading3')
                self._process_parameters(details.get('parameters', []))

                # 处理请求体
                if 'requestBody' in details:
                    self._add_paragraph("请求体:", style='Heading3')
                    content = details['requestBody'].get('content', {})
                    for content_type, schema in content.items():
                        self._add_paragraph(f"内容类型: {content_type}")
                        example = schema.get('example', '暂无示例')
                        self.doc.add_paragraph(str(example), style='Intense Quote')

                # 处理响应
                self._add_paragraph("响应:", style='Heading3')
                self._process_response(details.get('responses', {}))

                self.doc.add_page_break()

        # 保存文档
        self.doc.save(self.output_path)


# 使用示例
if __name__ == "__main__":
    converter = SwaggerToWord('api.json', 'output.docx')
    converter.generate_doc()