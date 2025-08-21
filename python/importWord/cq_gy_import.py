import base64
import os
import re
import uuid
import psycopg2
import logging
from docx import Document
from datetime import datetime


conn = psycopg2.connect(
    dbname="db",
    user="postgres",
    password="111111",
    host="192.168.1.10",
    port="9013"
)

current_time = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
log_filename = f"cqGY.py-{current_time}.log"
logging.basicConfig(filename=log_filename,level=logging.INFO)

# 处理文件内容
def delete_content(doc_path, output_path, keyword1='关键字1', keyword2='关键字2'):
    # 打开文档
    doc = Document(doc_path)

    # 查找关键字的第二次出现位置
    content_list = list(doc.element.body)
    first_count = 0
    second_count = 0
    first_index = -1
    second_index = -1

    for i, element in enumerate(content_list):
        if element.tag.endswith('p'):  # 段落
            if keyword1 in element.text:
                first_count += 1
                if first_count == 2 and first_index == -1:
                    first_index = i
            if keyword2 in element.text:
                second_count += 1
                if second_count == 2:
                    second_index = i
                    break  # 找到第二个关键字的第二次出现后就可以停止搜索

    if first_index == -1 or second_index == -1:
        logging.warning(f"{doc_path} 未找到 '{keyword1}' 或 '{keyword2}' 的第二次出现")
        return False

        # 删除第一个关键字第二次出现之前的所有内容
    for element in content_list[:first_index]:
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)

        # 删除第二个关键字第二次出现之后的所有内容
    for element in content_list[second_index:]:
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)

    #     # 删除所有剩余的内联图片
    # for paragraph in doc.paragraphs:
    #     for run in paragraph.runs:
    #         for inline in run._element.findall('.//wp:inline', namespaces=run._element.nsmap):
    #             inline.getparent().remove(inline)

    # 保存修改后的文档
    doc.save(output_path)
    logging.info(f"已保存修改后的文档到 {output_path}")
    return True

# 循环文件夹下的内容
def files(path, keyword='最新'):
    logging.info(f"开始执行-{path}")
    listdir = os.listdir(path)
    pattern = re.compile(keyword)
    for filename in listdir:
        if filename.endswith('.docx'):
            if pattern.search(filename):
                b = delete_content(os.path.join(path, filename), 'test/' + filename)

                if b is True:
                    db(filename)

    conn.close()




# 写入数据库
def db(filename):
    cur = conn.cursor()

    new_name = name(filename)
    base64_value = base('test/' + filename)
    uuid = chinese_to_uuid(filename)

    #file
    insert= "insert into file values (%s,%s,%s,%s,%s,%s,%s)"
    data= (uuid,base64_value,'2000-07-31 00:00:11.925808','9527','9527','1',filename)
    cur.execute(insert, data)
    conn.commit()

    #file_material
    select ="select * from file_material where status=1 and simple_name = %s "
    cur.execute(select,(new_name,))
    rows = cur.fetchall()
    if rows is not None:
        for row in rows:
            update = "update file_material set file_id = %s where id = %s "
            data = (uuid, row[0])
            cur.execute(update, data)
            conn.commit()

# 处理文件名称
def name(text):
    chinese_chars=re.findall(r'[\u4e00-\u9fa5]+', text)
    chinese_string = ''.join(chinese_chars)
    index=chinese_string.find('生产')
    if index != -1:
        result = chinese_string[:index]
        return result

# 获取文件base64的值
def base(file_path):
    doc = Document(file_path)

    with open(file_path, 'rb') as f:
        text_byte = f.read()

    base64_encoded=base64.b64encode(text_byte).decode('utf-8')
    return base64_encoded

def chinese_to_uuid(filename):
    bytes_text = filename.encode('utf-8')
    encode = base64.b64encode(bytes_text)
    uuid_text = uuid.uuid5(uuid.NAMESPACE_DNS, encode.decode('utf-8'))
    return f"{uuid_text}.docx"


if __name__ == '__main__':

    files(r'E:\word')